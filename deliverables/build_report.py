from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "API_Observability_Platform_Report.docx"

COLORS = {
    "ink": RGBColor(15, 23, 42),
    "muted": RGBColor(71, 85, 105),
    "blue": RGBColor(37, 99, 235),
    "green": RGBColor(5, 150, 105),
    "cyan": RGBColor(8, 145, 178),
    "light": "EFF6FF",
    "pale_green": "ECFDF5",
    "dark": "0F172A",
    "line": "CBD5E1",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CBD5E1", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill="0F172A"):
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            else:
                set_cell_shading(cell, "FFFFFF")


def add_title(document, title, subtitle):
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("API TESTING & MONITORING")
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = COLORS["blue"]
    run.bold = True

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    run.font.name = "Aptos Display"
    run.font.size = Pt(34)
    run.font.color.rgb = COLORS["ink"]
    run.bold = True

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(15)
    run.font.color.rgb = COLORS["muted"]

    table = document.add_table(rows=2, cols=4)
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(1.8)
    table.columns[2].width = Inches(1.8)
    table.columns[3].width = Inches(1.8)
    headers = ["Runtime", "Testing", "Observability", "Delivery"]
    values = ["Node.js + Express", "Postman, Newman, Jest, k6", "Prometheus + Grafana", "Docker Compose + CI"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=RGBColor(255, 255, 255), size=9)
        set_cell_shading(table.rows[0].cells[i], COLORS["dark"])
        set_cell_text(table.rows[1].cells[i], values[i], color=COLORS["ink"], size=8)
    style_table(table)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Prepared deliverable")
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["muted"]
    run.bold = True
    p = document.add_paragraph()
    run = p.add_run("Portfolio project report and implementation summary")
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["muted"]

    document.add_page_break()


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Aptos Display" if level == 1 else "Aptos"
        run.font.color.rgb = COLORS["ink"] if level == 1 else COLORS["blue"]
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["ink"]


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLORS["ink"]


def add_callout(document, title, text, fill="EFF6FF"):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "BFDBFE", "10")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = COLORS["blue"]
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = COLORS["ink"]


def build_document():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)

    add_title(
        doc,
        "API Observability Platform",
        "A full-stack API testing and monitoring system for synthetic checks, SLA reporting, Prometheus metrics, and Grafana dashboards.",
    )

    add_heading(doc, "Executive Summary")
    add_body(
        doc,
        "The API Observability Platform is a miniature production-style monitoring environment for API reliability work. It combines a realistic Express service, synthetic API tests, Prometheus metrics, Grafana visualization, load testing scripts, and Docker-based local orchestration.",
    )
    add_callout(
        doc,
        "Project outcome",
        "A single-command local stack that demonstrates API contract validation, SLA measurement, observability instrumentation, dashboard provisioning, and alert simulation.",
        "ECFDF5",
    )

    add_heading(doc, "System Architecture")
    add_body(
        doc,
        "The system follows a DevOps-friendly topology: clients and synthetic tests exercise the API, Express emits runtime metrics, Prometheus scrapes the metrics endpoint, and Grafana presents operational health, performance, and reliability signals.",
    )
    architecture = doc.add_table(rows=6, cols=3)
    widths = [1.55, 2.25, 3.15]
    for i, width in enumerate(widths):
        architecture.columns[i].width = Inches(width)
    rows = [
        ("Layer", "Component", "Responsibility"),
        ("API", "Node.js + Express", "Serves /health, /users, /posts, and /metrics with logging, timing, and error handling middleware."),
        ("Testing", "Postman, Newman, Jest", "Validates status codes, schemas, content, response-time SLA, and saves JSON run metrics."),
        ("Metrics", "prom-client + Prometheus", "Exposes counters, histograms, gauges, uptime, request volume, latency, and error counts."),
        ("Visualization", "Grafana", "Provisioned dashboard with health, performance, and reliability rows."),
        ("Delivery", "Docker Compose + GitHub Actions", "Runs the local stack and provides CI smoke checks for API and metrics readiness."),
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(architecture.rows[r].cells[c], value, bold=(r == 0), color=RGBColor(255, 255, 255) if r == 0 else COLORS["ink"], size=8.5)
    style_table(architecture)

    add_heading(doc, "Implementation Scope")
    scope = doc.add_table(rows=5, cols=3)
    rows = [
        ("Area", "Implemented assets", "Portfolio signal"),
        ("Express API", "Realistic user and post endpoints, health check, structured errors, request IDs, response timing", "Shows backend API design and operational middleware."),
        ("Synthetic testing", "Postman collection, Newman runner, Jest contract tests, generated SLA JSON", "Shows automated validation and test result normalization."),
        ("Observability", "Prometheus /metrics endpoint, scrape config, alert rules, Grafana dashboard JSON", "Shows SRE-style metrics thinking and dashboard design."),
        ("Operations", "Docker Compose, Dockerfile, k6 stress and spike scripts, CI workflow", "Shows local production simulation and automation discipline."),
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(scope.rows[r].cells[c], value, bold=(r == 0), color=RGBColor(255, 255, 255) if r == 0 else COLORS["ink"], size=8.4)
    style_table(scope, "1E293B")

    add_heading(doc, "Testing and SLA Strategy")
    add_body(
        doc,
        "The testing layer blends API contract checks with synthetic monitoring semantics. Postman tests use explicit SLA language, while Newman converts each run into normalized metrics that can be reviewed independently of the CLI output.",
    )
    add_bullets(
        doc,
        [
            "SLA target: Response time under 300ms for core synthetic requests.",
            "Coverage: GET /health, GET /users, GET /posts, and POST /posts.",
            "Validation: status codes, response schema, response content, response time, and generated metrics artifact shape.",
            "Artifacts: metrics/newman/latest-run.json and metrics/newman/latest-summary.json.",
        ],
    )

    add_heading(doc, "Dashboard Design")
    dashboard = doc.add_table(rows=4, cols=3)
    rows = [
        ("Dashboard row", "Panels", "Operational question answered"),
        ("System Health", "Uptime, average latency, total requests, SLA compliance", "Is the service currently healthy and meeting its stated SLA?"),
        ("Performance", "Response time trend, endpoint latency comparison, requests per second", "Where is traffic or latency shifting over time?"),
        ("Reliability", "Error rate, failed requests, recent failures table", "What is failing, how often, and on which route/status combination?"),
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(dashboard.rows[r].cells[c], value, bold=(r == 0), color=RGBColor(255, 255, 255) if r == 0 else COLORS["ink"], size=8.7)
    style_table(dashboard, "075985")

    add_heading(doc, "Local Runbook")
    add_body(doc, "The project is designed for minimal local setup.")
    runbook = doc.add_table(rows=5, cols=2)
    rows = [
        ("Command", "Purpose"),
        ("npm install", "Install Node dependencies."),
        ("docker compose up --build", "Start API, Prometheus, and Grafana."),
        ("npm test", "Run Jest API and metrics artifact tests."),
        ("npm run test:postman:metrics", "Run Newman synthetic SLA suite and save JSON metrics."),
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(runbook.rows[r].cells[c], value, bold=(r == 0), color=RGBColor(255, 255, 255) if r == 0 else COLORS["ink"], size=8.8)
    style_table(runbook, "166534")

    add_heading(doc, "Validation Summary")
    add_bullets(
        doc,
        [
            "Jest suite passed with 7 tests covering API contracts, error handling, Prometheus output, and example metrics format.",
            "Newman synthetic suite passed with 4 requests, 16 assertions, and 0 failures.",
            "Generated Newman summary reported 100% SLA compliance in the latest local run.",
            "Docker Compose configuration was parsed successfully with API, Prometheus, Grafana, networks, and persistent volumes.",
        ],
    )

    add_heading(doc, "Recommended Next Enhancements")
    add_bullets(
        doc,
        [
            "Persist synthetic run history into a lightweight datastore for longitudinal SLA reporting.",
            "Add Alertmanager routing to demonstrate notification flows.",
            "Publish dashboard screenshots after running the stack with k6 traffic.",
            "Add OpenAPI documentation and contract drift checks.",
        ],
    )

    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("API Observability Platform | Project report")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
